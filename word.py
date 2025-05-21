import os
from datetime import datetime
from docx import Document

# Summary folder
summary_dir = "./summaries"

# Word output folder
word_dir = "./word"
os.makedirs(word_dir, exist_ok=True)

# Create Word document
doc = Document()
doc.add_heading("AI News Summaries", level=0)

# Read all summary txt files
for filename in sorted(os.listdir(summary_dir)):
    if filename.endswith("_summary.txt"):
        filepath = os.path.join(summary_dir, filename)
        with open(filepath, "r", encoding="utf-8") as f:
            content = f.read()

        # Parse Title and Summary sections
        try:
            title, summary = content.split("Summary:\n", 1)
        except ValueError:
            continue 

        title = title.strip().replace("Title: ", "")
        summary = summary.strip()

        # Add title and summary to Word document
        doc.add_heading(title, level=1)
        for line in summary.splitlines():
            if line.strip():
                doc.add_paragraph(line.strip(), style="List Bullet")

# Save the Word file with today's date in the filename
today_str = datetime.now().strftime("%Y-%m-%d")
output_path = os.path.join(word_dir, f"AINEWS_{today_str}.docx")
doc.save(output_path)
print(f"Summary completed. Output file: {output_path}")
